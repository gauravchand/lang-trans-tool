import gradio as gr
import sys
import pkg_resources
import tempfile
import os
from pathlib import Path

def check_dependencies():
    required_packages = {
        'gradio': ['gradio'],
        'transformers': ['transformers'],
        'python-docx': ['python-docx', 'python_docx', 'docx'],
        'PyPDF2': ['PyPDF2', 'pypdf2', 'pypdf'],
        'torch': ['torch'],
        'sentencepiece': ['sentencepiece'],
        'tf-keras': ['tf-keras']  # Added tf-keras as a required package
    }
    
    installed = {pkg.key.lower() for pkg in pkg_resources.working_set}
    missing = []
    
    for package, variations in required_packages.items():
        if not any(variation.lower() in installed for variation in variations):
            missing.append(package)
    
    if missing:
        print("Missing required packages. Please install:")
        for pkg in missing:
            print(f"pip install {pkg}")
        sys.exit(1)

# Check dependencies before importing
check_dependencies()

import torch
from transformers import pipeline
import docx
import PyPDF2
import io

class DocumentTranslator:
    def __init__(self):
        try:
            # Initialize translation model with PyTorch backend explicitly
            self.translator = pipeline(
                "translation",
                model="Helsinki-NLP/opus-mt-en-ROMANCE",
                framework="pt"  # Explicitly specify PyTorch as the backend
            )
            
            # Supported languages
            self.languages = {
                "English": "en",
                "French": "fr",
                "Spanish": "es",
                "Portuguese": "pt",
                "Italian": "it"
            }
        except Exception as e:
            print(f"Error initializing translator: {str(e)}")
            print("Please make sure all required packages are installed:")
            print("pip install transformers torch sentencepiece python-docx PyPDF2 gradio tf-keras")
            raise

    def extract_text_from_docx(self, file):
        doc = docx.Document(file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    def extract_text_from_pdf(self, file):
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text())
        return "\n".join(text)

    def create_translated_docx(self, original_text, translated_text, output_filename):
        doc = docx.Document()
        paragraphs = translated_text.split("\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        doc.save(output_filename)
        return output_filename

    def translate_document(self, file, source_lang, target_lang):
        try:
            # Create temporary directory for output
            temp_dir = tempfile.mkdtemp()
            output_filename = os.path.join(temp_dir, "translated_document.docx")

            # Extract text based on file type
            if file.name.endswith('.docx'):
                text = self.extract_text_from_docx(file)
            elif file.name.endswith('.pdf'):
                text = self.extract_text_from_pdf(file)
            else:
                return None, "Unsupported file format. Please use .docx or .pdf"

            # Split text into chunks to handle long documents
            chunk_size = 500
            chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
            
            # Translate chunks
            translated_chunks = []
            for chunk in chunks:
                translation = self.translator(chunk)[0]['translation_text']
                translated_chunks.append(translation)
            
            translated_text = " ".join(translated_chunks)
            
            # Create new document with translation
            output_file = self.create_translated_docx(text, translated_text, output_filename)
            
            return output_file, "Translation completed successfully!"
            
        except Exception as e:
            return None, f"Error during translation: {str(e)}"

def create_translation_interface():
    try:
        translator = DocumentTranslator()
        
        def translate_file(file, source_lang, target_lang):
            if file is None:
                return None, "Please upload a file"
            return translator.translate_document(file, source_lang, target_lang)

        iface = gr.Interface(
            fn=translate_file,
            inputs=[
                gr.File(label="Upload Document (.docx or .pdf)"),
                gr.Dropdown(choices=list(translator.languages.keys()), label="Source Language"),
                gr.Dropdown(choices=list(translator.languages.keys()), label="Target Language")
            ],
            outputs=[
                gr.File(label="Download Translated Document"),
                gr.Textbox(label="Status")
            ],
            title="Document Translation System",
            description="Upload a document (.docx or .pdf) and select source and target languages for translation.",
            theme="default"
        )
        
        return iface
    except Exception as e:
        print(f"Error creating interface: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    print("Initializing translation system...")
    print("Checking dependencies...")
    check_dependencies()
    print("Starting Gradio interface...")
    iface = create_translation_interface()
    iface.launch(share=True)